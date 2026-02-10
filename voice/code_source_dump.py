import os
from docx import Document
from docx.shared import Pt


import os
from docx import Document
from docx.shared import Pt
import re


# XML-safe text filter
def make_xml_safe(text: str) -> str:
    # Remove NULL bytes and illegal control chars
    return re.sub(
        r"[\x00-\x08\x0B\x0C\x0E-\x1F]",
        "",
        text
    )


def extract_code_files_to_word(folder_path: str, output_docx: str):
    print(f"[DEBUG] Starting extraction from folder: {folder_path}")
    print("[DEBUG] Mode: ALL extensions (text-safe only)")

    if not os.path.isdir(folder_path):
        print("[ERROR] Provided path is not a directory")
        return

    output_docx = os.path.abspath(output_docx)

    document = Document()
    document.add_heading("Source Code Dump (All Extensions)", level=1)

    file_count = 0
    skipped_files = 0

    for root, _, files in os.walk(folder_path):
        for file in files:
            file_path = os.path.abspath(os.path.join(root, file))

            # 🚫 Skip the output file itself
            if file_path == output_docx:
                print(f"[SKIP] Output file: {file_path}")
                continue

            print(f"[DEBUG] Attempting file: {file_path}")

            try:
                with open(file_path, "rb") as f:
                    raw = f.read()

                # 🚫 Binary detection (NULL byte)
                if b"\x00" in raw:
                    skipped_files += 1
                    print(f"[SKIP] Binary file detected: {file_path}")
                    continue

                # Decode text
                try:
                    content = raw.decode("utf-8")
                except UnicodeDecodeError:
                    content = raw.decode("latin-1")

                # XML safety
                content = make_xml_safe(content)

            except Exception as e:
                skipped_files += 1
                print(f"[SKIP] Error reading file: {file_path} | {e}")
                continue

            file_count += 1
            ext = os.path.splitext(file)[1].lower()

            document.add_page_break()
            document.add_heading(
                f"File: {file_path} ({ext if ext else 'no extension'})",
                level=2
            )

            para = document.add_paragraph()
            run = para.add_run(content)
            run.font.name = "Courier New"
            run.font.size = Pt(9)

    document.save(output_docx)

    print("[DEBUG] Extraction complete")
    print(f"[DEBUG] Files added: {file_count}")
    print(f"[DEBUG] Files skipped: {skipped_files}")
    print(f"[DEBUG] Word document saved as: {output_docx}")


# ===============================
# USAGE
# ===============================
if __name__ == "__main__":
    SOURCE_FOLDER = r"C:\Users\Ridhv\Desktop\websites\aiass\voice"
    OUTPUT_FILE = "source_code_dump.docx"

    extract_code_files_to_word(SOURCE_FOLDER, OUTPUT_FILE)